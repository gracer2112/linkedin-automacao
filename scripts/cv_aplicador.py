# cv_aplicador.py

import sys
from dotenv import load_dotenv
from docx import Document
import os
import json
import logging
import subprocess
import time
import re # Para regex na interpretação da resposta (se necessário, mas não deve mais ser neste script)

# ================= LOGGING SETUP =================
loglevel = os.environ.get("MY_LOG_LEVEL", "INFO").upper()
logger = logging.getLogger(__name__)
logger.setLevel(loglevel)

if not logger.handlers:
    handler = logging.StreamHandler(sys.stderr)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

# ================= FUNÇÕES DO SCRIPT =================

def stdin_has_data():
    """Verifica se há dados na entrada padrão (stdin)."""
    import select
    if sys.stdin.isatty():
        return False
    rlist, _, _ = select.select([sys.stdin], [], [], 0)
    return bool(rlist)

def criar_diretorio_vaga(base_dir, codigo_vaga):
    """Cria um diretório para a vaga específica."""
    dir_destino = os.path.join(base_dir, str(codigo_vaga))
    os.makedirs(dir_destino, exist_ok=True)
    logger.info(f"Diretório criado/verificado: '{dir_destino}' para a vaga '{codigo_vaga}'")
    return dir_destino

def gerar_pdf_linux(docx_path, output_dir):
    """Converte um arquivo DOCX para PDF usando LibreOffice (para Linux)."""
    logger.info(f"Iniciando conversão de '{os.path.basename(docx_path)}' para PDF...")
    
    caminho_docx_abs = os.path.abspath(docx_path)
    diretorio_saida_abs = os.path.abspath(output_dir)

    if not os.path.exists(caminho_docx_abs):
        logger.error(f"Arquivo DOCX não encontrado em: {caminho_docx_abs}")
        return False

    if not os.path.exists(diretorio_saida_abs):
        logger.error(f"Diretório de saída não encontrado em: {diretorio_saida_abs}")
        return False
    
    cmd_list = [    
        "libreoffice",
        "--headless",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--convert-to", "pdf",
        "--outdir", diretorio_saida_abs,
        caminho_docx_abs
    ]
    
    logger.info(f"Executando conversão para PDF: {' '.join(cmd_list)}")

    try:
        resultado = subprocess.run(
            cmd_list, 
            check=True, 
            capture_output=True, 
            text=True
        )
        logger.info(f"LibreOffice stdout: {resultado.stdout}")
        logger.info(f"Conversão para PDF bem-sucedida para {os.path.basename(docx_path)}")
        # Retorna o caminho completo do PDF gerado
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        return os.path.join(diretorio_saida_abs, pdf_filename)

    except FileNotFoundError:
        logger.error("Erro: O comando 'libreoffice' não foi encontrado. Verifique se está no PATH do sistema.")
        logger.error("Dica: No Ubuntu/Debian, instale com 'sudo apt-get install libreoffice'.")
        return False
    except subprocess.CalledProcessError as e:
        logger.error("Erro ao executar o comando do LibreOffice.")
        logger.error(f"Código de Retorno: {e.returncode}")
        logger.error(f"Stdout: {e.stdout}")
        logger.error(f"Stderr: {e.stderr}")
        return False
    except Exception as e:
        logger.error(f"Erro inesperado durante a conversão para PDF: {e}")
        return False

def extrair_texto_docx(caminho_arquivo):
    """Extrai texto de um arquivo DOCX (função auxiliar, pode ser redundante se o CV já foi lido antes)."""
    try:
        doc = Document(caminho_arquivo)
        texto = []
        for p in doc.paragraphs:
            texto.append(p.text)
        return "\n".join(texto)
    except Exception as e:
        logger.error(f"Erro ao extrair texto do DOCX '{caminho_arquivo}': {e}")
        # Não sys.exit(1) aqui, pois pode ser que o arquivo esteja corrompido, mas queremos continuar
        # processando outras vagas, se houver.
        return None

def substituir_texto_docx(caminho_arquivo_original, substituicoes):
    """
    Realiza as substituições no documento DOCX, preservando a formatação.
    Retorna o objeto Document modificado ou None em caso de erro.
    """
    try:
        document = Document(caminho_arquivo_original)
        logger.info(f"Documento '{os.path.basename(caminho_arquivo_original)}' carregado para substituição.")

        # Combina o processamento para parágrafos e tabelas
        todos_paragrafos = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    todos_paragrafos.extend(cell.paragraphs)

        for p in todos_paragrafos:
            for sub in substituicoes:
                original = str(sub.get("original", ""))
                substituto = str(sub.get("substituto", ""))

                if not original or not substituto:
                    logger.warning(f"Sugestão de substituição inválida: {sub}. Ignorando.")
                    continue

                # 1. Junta todas as 'runs' para encontrar o texto completo no parágrafo
                texto_completo_paragrafo = ''.join(run.text for run in p.runs)

                # 2. Se o texto a ser substituído existe no parágrafo completo...
                if original in texto_completo_paragrafo:
                    logger.info(f"Encontrado '{original}' no parágrafo. Realizando substituição.")
                    
                    # Para substituir preservando formatação, precisamos reconstruir o parágrafo.
                    # A maneira mais simples (e que o código original já tentava) é recriar o parágrafo
                    # ou manipular as runs diretamente. A abordagem de Python-Docx para isso é complexa
                    # se o termo original estiver quebrado entre runs.
                    # Uma alternativa pragmática é substituir o texto completo e re-aplicar uma formatação padrão
                    # ou tentar copiar a formatação da primeira run.
                    
                    # Abordagem simplificada (pode perder formatação granular em certos casos complexos):
                    # Se o parágrafo tiver apenas uma run, ou se for simples, essa abordagem funciona bem.
                    if len(p.runs) == 1 and original in p.runs[0].text:
                        p.runs[0].text = p.runs[0].text.replace(original, substituto)
                        logger.info(f"Substituição direta em run única: '{original}' por '{substituto}'.")
                    else:
                        # Para casos mais complexos, onde o texto está em várias runs,
                        # é mais seguro substituir o texto completo e, em seguida,
                        # limpar as runs existentes e adicionar uma nova run com o texto substituído.
                        # Isso garante que a substituição ocorra, mas pode redefinir o estilo
                        # de todo o parágrafo para o estilo padrão, a menos que seja re-aplicado.

                        # Salva o estilo do parágrafo antes de limpar
                        paragrafo_style = p.style

                        # Limpa as runs existentes (remove todo o texto)
                        for run in p.runs:
                            run.text = ""
                        
                        # Adiciona uma nova run com o texto substituído
                        new_text = texto_completo_paragrafo.replace(original, substituto)
                        p.add_run(new_text)
                        
                        # Tenta restaurar o estilo do parágrafo
                        p.style = paragrafo_style
                        
                        logger.info(f"Substituição complexa em múltiplas runs: '{original}' por '{substituto}'.")
        
        return document

    except Exception as e:
        logger.error(f"Erro ao processar o documento '{caminho_arquivo_original}' para substituição: {e}")
        return None

# ================= FLUXO PRINCIPAL =================

def main():
    load_dotenv()
    processing_results = []

    try:
        # Carregar informações das configurações
        config_path = os.environ.get('CONFIG_JSON_PATH', 'configs/linkedin.json')
        with open(config_path, "r", encoding="utf-8") as file:
            linkedin_config = json.load(file)
    except Exception as e:
        logger.error(f"Erro ao carregar configuração: {e}")
        sys.exit(1)

    diretorio_cv_original = linkedin_config["input_file_cv"]
    output_base_dir = linkedin_config['output_dir']

    # Lê o JSON consolidado de sugestões (entrada via stdin ou arquivo)
    try:
        if stdin_has_data():
            entrada_json = sys.stdin.read().strip()
            logger.info("Lido JSON de sugestões do stdin.")
        elif len(sys.argv) > 1:
            with open(sys.argv[1]) as f:
                entrada_json = f.read()
            logger.info(f"Lido JSON de sugestões do arquivo {sys.argv[1]}.")
        else:
            logger.info("Nenhuma entrada de sugestões fornecida! Use argumento de arquivo ou STDIN.")
            sys.exit(1)
        
        all_vaga_suggestions = json.loads(entrada_json)
        if not isinstance(all_vaga_suggestions, list):
            # Se for um único objeto, transforme em lista para processamento
            all_vaga_suggestions = [all_vaga_suggestions]

    except Exception as e:
        logger.error(f"Erro ao ler entrada de sugestões: {e}")
        sys.exit(1)
    
    if not os.path.exists(diretorio_cv_original):
        logger.error(f"Arquivo do CV original não encontrado: {diretorio_cv_original}")
        sys.exit(1)

    # Itera sobre cada vaga e suas sugestões
    for vaga_data in all_vaga_suggestions:
        codigo_vaga = vaga_data.get("codigo", "SEM_CODIGO")
        sugestoes = vaga_data.get("sugestoes", [])
        referencia = vaga_data.get("referencia", {})

        logger.info(f"\nProcessando documentos para a vaga: {codigo_vaga}")
        
        dir_vaga = criar_diretorio_vaga(output_base_dir, codigo_vaga)
        novo_docx_path = os.path.join(dir_vaga, f"CV_Modificado_{codigo_vaga}.docx")

        if not sugestoes:
            logger.warning(f"Nenhuma sugestão de otimização para a vaga {codigo_vaga}. Pulando modificação do DOCX.")
            # Opcional: Copiar o CV original para o diretório da vaga mesmo sem modificações
            # import shutil
            # shutil.copy(diretorio_cv_original, novo_docx_path)
            # logger.info(f"Copiado CV original para {novo_docx_path} (sem modificações).")
            processing_results.append({
                "codigo": codigo_vaga,
                "status": "Nenhuma sugestão, DOCX original não modificado",
                "output_dir": dir_vaga,
                "referencia": referencia
            })
            continue # Passa para a próxima vaga

        # Aplicar sugestões e salvar novo DOCX
        sucesso_docx_obj = substituir_texto_docx(diretorio_cv_original, sugestoes)

        if sucesso_docx_obj:
            try:
                sucesso_docx_obj.save(novo_docx_path)
                logger.info(f"Documento modificado salvo como '{os.path.basename(novo_docx_path)}'.")
           
                # Gerar PDF
                caminho_pdf_gerado = gerar_pdf_linux(novo_docx_path, dir_vaga)
                if caminho_pdf_gerado:
                    logger.info(f"PDF gerado com sucesso em: {caminho_pdf_gerado}")
                    processing_results.append({
                        "codigo": codigo_vaga,
                        "status": "Sucesso",
                        "output_docx": novo_docx_path,
                        "output_pdf": caminho_pdf_gerado,
                        "output_dir": dir_vaga,
                        "sugestoes_aplicadas": sugestoes,
                        "referencia": referencia
                    })
                else:
                    logger.error(f"Falha ao gerar PDF para a vaga {codigo_vaga}.")
                    processing_results.append({
                        "codigo": codigo_vaga,
                        "status": "Falha ao gerar PDF",
                        "output_docx": novo_docx_path,
                        "output_dir": dir_vaga,
                        "sugestoes_aplicadas": sugestoes,
                        "referencia": referencia
                    })
            except Exception as e:
                logger.error(f"Falha ao tentar salvar o DOCX modificado para a vaga {codigo_vaga}: {e}")
                processing_results.append({
                    "codigo": codigo_vaga,
                    "status": "Falha ao salvar DOCX modificado",
                    "output_dir": dir_vaga,
                    "erro": str(e),
                    "referencia": referencia
                })
        else:
            logger.error(f"Falha ao modificar o arquivo DOCX para a vaga {codigo_vaga}. Nenhum arquivo salvo.")
            processing_results.append({
                "codigo": codigo_vaga,
                "status": "Falha ao modificar DOCX",
                "output_dir": dir_vaga,
                "referencia": referencia
            })
        
        time.sleep(0.5) # Pequeno atraso para evitar sobrecarga de I/O ou CPU

    # Imprime o JSON de resultados de processamento para o stdout
    print(json.dumps(processing_results, ensure_ascii=False, indent=2))
    logger.info("Processo de aplicação de sugestões e geração de documentos finalizado.")

if __name__ == "__main__":
    main()