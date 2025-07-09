#cv_otimizado.py
from dotenv import load_dotenv
from docx import Document
import os
import json
from google.auth import default
import google.generativeai as genai
import ast
import re
import subprocess
import sys
import logging
import time

# ================= LOGGING SETUP =================
# Configuração mais robusta e explícita do logging
# Define o nível de log a partir de uma variável de ambiente, com 'INFO' como padrão.
loglevel = os.environ.get("MY_LOG_LEVEL", "INFO").upper()

# Cria um logger específico para este módulo em vez de usar o root logger. Boa prática.
logger = logging.getLogger(__name__)
logger.setLevel(loglevel)

# Evita adicionar múltiplos handlers se o script for importado em outro lugar.
if not logger.handlers:
    # Direciona explicitamente os logs para stderr, que é o que você quer capturar com '2>'
    handler = logging.StreamHandler(sys.stderr)
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

# ===================================================

load_dotenv()

def carrega_chave():
    gemini_api_key = os.getenv('GEMINI_API_KEY')
    if not gemini_api_key:
        logger.error("Chave de API do Google não encontrada (GEMINI_API_KEY).")
        sys.exit(1)
        
    genai.configure(api_key=gemini_api_key)
    return genai

def stdin_has_data():
    import select
    if sys.stdin.isatty():
        return False
    rlist, _, _ = select.select([sys.stdin], [], [], 0)
    return bool(rlist)

def criar_diretorio_vaga(base_dir, codigo_vaga):
    dir_destino = os.path.join(base_dir, str(codigo_vaga))
    os.makedirs(dir_destino, exist_ok=True)
    logger.info(f"Iniciando processamento para a vaga: '{codigo_vaga}'")
    return dir_destino

# def gerar_pdf(docx_path, pasta_dest):
#     pdf_path = os.path.join(pasta_dest, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
#     convert(docx_path, pdf_path)
#     return pdf_path

def gerar_pdf_linux(docx_path, output_dir):
    """Converte um arquivo DOCX para PDF usando LibreOffice."""
    logger.info(f"Iniciando conversão de '{os.path.basename(docx_path)}' para PDF...")
    # --- PONTO DA CORREÇÃO ---
    # Crie um ambiente customizado para o subprocesso
    #env_customizado = os.environ.copy()
    #env_customizado['HOME'] = output_dir  # Força o LibreOffice a usar o diretório de saída como "casa"

    # Garante que os caminhos sejam absolutos para evitar problemas de diretório de trabalho
    caminho_docx_abs = os.path.abspath(docx_path)
    diretorio_saida_abs = os.path.abspath(output_dir)

    if not os.path.exists(caminho_docx_abs):
        logger.error(f"Arquivo DOCX não encontrado em: {caminho_docx_abs}")
        return False

    if not os.path.exists(diretorio_saida_abs):
        logger.error(f"Diretório de saída não encontrado em: {diretorio_saida_abs}")
        return False
    
    cmd_list = [    
        "libreoffice",
        "--headless",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--convert-to", "pdf",
        "--outdir", diretorio_saida_abs,
        caminho_docx_abs
    ]
    #subprocess.run(cmd_list,check=True,env=env_customizado)
    logger.info(f"Executando conversão para PDF: {' '.join(cmd_list)}")

    try:
        # capture_output=True para pegar o stdout/stderr
        # text=True para decodificar como texto
        # check=True para lançar uma exceção se o comando retornar um erro
        resultado = subprocess.run(
            cmd_list, 
            check=True, 
            capture_output=True, 
            text=True
        )
        logger.info(f"LibreOffice stdout: {resultado.stdout}")
        logger.info(f"Conversão para PDF bem-sucedida para {output_dir}")
        return True

    except FileNotFoundError:
        logger.error("Erro: O comando 'libreoffice' não foi encontrado. Verifique se está no PATH do sistema.")
        return False
    except subprocess.CalledProcessError as e:
        # ESTA É A PARTE MAIS IMPORTANTE PARA DEBUG!
        logger.error("Erro ao executar o comando do LibreOffice.")
        logger.error(f"Código de Retorno: {e.returncode}")
        logger.error(f"Stdout: {e.stdout}")
        logger.error(f"Stderr: {e.stderr}") # O erro de permissão provavelmente aparecerá aqui!
        return False

def interpretar_resposta_ia(resposta_texto):

    logging.warning("Resposta da IA não parece ser um JSON. Validando entrada.")

    # Remove marcações de código e espaços em branco desnecessários
    cleaned_text = re.sub(r'```json\s*|```', '', resposta_texto).strip()

    # Se a resposta estiver vazia após a limpeza, retorna lista vazia.
    if not cleaned_text:
        logging.warning("Resposta da IA estava vazia após a limpeza.")
        return []

    try:
        # Tenta converter a string para um objeto Python
        dados_convertidos = ast.literal_eval(cleaned_text)

        # Caso 1: A resposta já é uma lista (formato ideal)
        if isinstance(dados_convertidos, list):
            logging.info("Resposta da IA interpretada como uma lista de sugestões.")
            return dados_convertidos

        # Caso 2: A resposta é um dicionário contendo a chave 'sugestoes'
        elif isinstance(dados_convertidos, dict):
            if 'sugestoes' in dados_convertidos and isinstance(dados_convertidos['sugestoes'], list):
                logging.info("Resposta da IA interpretada como um dicionário. Extraindo a lista 'sugestoes'.")
                return dados_convertidos['sugestoes']
            else:
                logging.error(f"Dicionário JSON não contém a chave 'sugestoes' ou o valor não é uma lista. Resposta recebida: {cleaned_text}")
                return []
        
        # Caso 3: Formato inesperado (nem lista, nem dicionário)
        else:
            logging.error(f"Resposta da IA não é uma lista ou dicionário válido. Tipo recebido: {type(dados_convertidos)}")
            return []
        
    except (ValueError, SyntaxError, TypeError) as e:
        # Erro de parsing: A string não é uma estrutura Python válida
        logging.error(f"Erro de parsing ao interpretar a resposta da IA: {e}. Resposta bruta recebida: {resposta_texto}")
        return [] # Retorna lista vazia para não quebrar o fluxo principal
    except Exception as e:
        # Pega qualquer outro erro inesperado
        logging.error(f"Erro genérico inesperado ao interpretar resposta: {e}. Resposta bruta recebida: {resposta_texto}")
        return []

# 1. Função para gerar sugestões com Vertex IA / Gemini (Google GenAI)
def sugerir_substituicoes(genai, texto_cv, requisitos_vaga, model="gemini-1.5-flash"):
#def sugerir_substituicoes(genai, texto_cv, requisitos_vaga, model="gemini-2.5-flash-preview-05-20"):
    # Autentica usando application default credentials
    creds, _ = default()
    response = None  # Inicializa a variável como None

    prompt = f"""
Você é um especialista em RH e otimização de currículos. Analise o currículo e os requisitos da vaga abaixo.
Sugira substituições de termos no currículo para que ele se alinhe melhor aos requisitos da vaga.
Concentre-se em trocar jargões internos ou termos menos comuns por palavras-chave presentes na descrição da vaga ou mais reconhecidas no mercado.

**Formato da Resposta:**
Sua resposta deve ser ESTRITAMENTE um array JSON válido (uma lista de objetos JSON). Cada objeto deve conter duas chaves: "original" (o termo do CV) e "substituto" (a sugestão).
NÃO inclua nenhuma explicação, texto introdutório ou formatação extra. Apenas o array JSON.
Seja sucinto e prático, não sugira mais do que 4 itens!

### Exemplo de Saída Esperada:
[
  {{"original": "Termo Antigo 1", "substituto": "Termo Novo 1"}},
  {{"original": "Termo Antigo 2", "substituto": "Termo Novo 2"}}
]

-----------------
{texto_cv}
-----------------
Compare com estes requisitos da vaga:
{requisitos_vaga}
-----------------

"""
    try:
        # Chamada para gerar o conteúdo
        logger.info(f"Enviando o seguinte prompt para a IA:\n{prompt}")
        modelo = genai.GenerativeModel(model)
        config =  genai.types.GenerationConfig(temperature=0.7,response_mime_type="application/json")
        response = modelo.generate_content(
            prompt,  # Usando a nova API com Part
            generation_config=config
        )
        logger.info(f"DEBUG: repr(response.text) antes de interpretar:\n {repr(response.text)}")    
    except Exception as e:
        logger.error(f"Ocorreu um erro ao chamar a API: {e}")
        logger.exception("Erro ao interpretar resposta IA")
        # Verifica se a resposta contém texto
    

    sugestoes_ia = interpretar_resposta_ia(response.text)
    logger.info(f"DEBUG: Sugestões interpretadas da IA: {sugestoes_ia}")
    return sugestoes_ia

    
# 2. Função para extrair texto de um DOCX (simplificada)
def extrair_texto_docx(caminho_arquivo):
    doc = Document(caminho_arquivo)
    texto = []
    for p in doc.paragraphs:
        texto.append(p.text)
    return "\n".join(texto)

def substituir_texto_docx(caminho_arquivo_original, substituicoes):
    try:
        document = Document(caminho_arquivo_original)
        logger.info(f"Documento '{os.path.basename(caminho_arquivo_original)}' carregado.")

        # Combina o processamento para parágrafos e tabelas para evitar duplicação
        todos_paragrafos = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    todos_paragrafos.extend(cell.paragraphs)

        for p in todos_paragrafos:
            for sub in substituicoes:
                original = str(sub["original"]) # Garante que é string
                substituto = str(sub["substituto"]) # Garante que é string

                # 1. Junta todas as 'runs' para encontrar o texto completo no parágrafo
                texto_completo_paragrafo = ''.join(run.text for run in p.runs)

                # 2. Se o texto a ser substituído existe no parágrafo completo...
                if original in texto_completo_paragrafo:
                    logger.info(f"Encontrado '{original}' no parágrafo. Realizando substituição complexa.")
                    
                    # 3. Lógica de substituição que preserva a formatação
                    for i, run in enumerate(p.runs):
                        if original in run.text:
                            # Caso simples: o texto está contido inteiramente em uma run
                            run.text = run.text.replace(original, substituto)
                            break # Sai do loop de runs após substituir
                        
                        # Início do caso complexo: o texto está espalhado por várias runs
                        # Procura o início do texto original em várias runs
                        temp_text = ''.join(r.text for r in p.runs[i:])
                        if temp_text.startswith(original):
                            # Encontrou o início. Agora apaga as runs necessárias.
                            # Guarda a formatação da primeira run
                            formato_run = run.style 
                            
                            runs_a_remover = 0
                            texto_acumulado = ""
                            for j in range(i, len(p.runs)):
                                texto_acumulado += p.runs[j].text
                                runs_a_remover += 1
                                if original in texto_acumulado:
                                    # Texto final encontrado, remove as runs e insere a nova
                                    # Limpa o texto das runs que serão removidas
                                    for k in range(i, i + runs_a_remover):
                                        p.runs[k].text = ""
                                    
                                    # Adiciona o texto substituído na primeira run da sequência
                                    p.runs[i].text = texto_completo_paragrafo.replace(original, substituto)
                                    #p.runs[i].style = formato_run # Re-aplica estilo se necessário
                                    
                                    logger.info(f"Substituído '{original}' por '{substituto}'.")
                                    break # Sai do loop
                            break # Sai do loop de substituições para este parágrafo
        
        return document

    except Exception as e:
        logger.error(f"Erro ao processar o documento {caminho_arquivo_original}: {e}")
        return None

def processar_vaga(genai, vaga_dict, output_dir,caminho_cv):
    # Montar string dos requisitos
    # Prioriza os dados dentro de "analise"
    analise = vaga_dict.get("analise", {})
    referencia = vaga_dict.get("referencia", {})

    logger.info(f"Iniciando processamento para a vaga: {referencia.get('Code', 'N/A')}")

    linhas_requisitos = []

    campos_lista = [
        "requisitos_obrigatorios",
        "requisitos_desejaveis",
        "soft_skills",
        "hard_skills"
    ]

    for campo in campos_lista:
        if campo in analise and isinstance(analise[campo], list):
            linhas_requisitos.append(f"{campo.replace('_',' ').capitalize()}: {', '.join(analise[campo])}")
    for campo, valor in analise.items():
        if campo not in campos_lista and not isinstance(valor, (list, dict)):
            linhas_requisitos.append(f"{campo}: {valor}")
    requisitos_texto = "\n".join(linhas_requisitos)

    codigo_vaga = str(int(referencia.get('Code', "SEM_CODIGO")))
    logger.info(f"Processando vaga..., {codigo_vaga}", file=sys.stderr)
    dir_vaga = criar_diretorio_vaga(output_dir, codigo_vaga)
    novo_docx_path = os.path.join(dir_vaga, f"CV_Modificado_{codigo_vaga}.docx")

    # Processamento
    texto_cv = extrair_texto_docx(caminho_cv)
    logger.info("Solicitando sugestões IA...")
    sugestoes_ia = sugerir_substituicoes(genai, texto_cv, requisitos_texto)
    logger.info(f"Sugestões:, {sugestoes_ia}")

    # Substituições manuais, se necessário
    sugestoes_ia.append({"original":"Brasileira, solteira, 52 anos, sem filhos",
                        "substituto":"Brasileira, solteira, 53 anos, sem filhos"})
    logging.info(f"Adicionada correção manual da idade. Total de sugestões: {len(sugestoes_ia)}")

    # Crie o nome do arquivo usando o código da vaga (garanta que não há caracteres problemáticos!)
    nome_arquivo = f"resultado_{codigo_vaga}.json"
    caminho_arquivo = os.path.join(dir_vaga, nome_arquivo)
    with open(caminho_arquivo, "w", encoding="utf-8") as f:
        json.dump(sugestoes_ia, f, ensure_ascii=False, indent=4)
    logger.info(f"Sugestões salvas em: {caminho_arquivo}")

    #sucesso_docx = substituir_e_salvar_docx(caminho_cv, novo_docx_path, sugestoes_ia)
    sucesso_docx = substituir_texto_docx(caminho_cv, sugestoes_ia)

    if sucesso_docx:
        try:
            sucesso_docx.save(novo_docx_path)
            logger.info(f"Documento modificado salvo como '{os.path.basename(novo_docx_path)}'.")
       
            # PDF se desejar, escolha o método de acordo com o sistema
            caminho_pdf_gerado = gerar_pdf_linux(novo_docx_path, dir_vaga)
            if caminho_pdf_gerado:
                logger.info(f"PDF gerado com sucesso em: {caminho_pdf_gerado}")
        except Exception as e:
            logger.error(f"Falha ao tentar gerar o DOCX para a vaga {codigo_vaga}: {e}")
    else:
        logger.error(f"Falha ao modificar o arquivo DOCX para a vaga {codigo_vaga}. Nenhum arquivo salvo")
    
    time.sleep(1)  # Respeita limites de API

    return {"codigo" : codigo_vaga,
            "sugestoes": sugestoes_ia
            }

# --------------- FLUXO PRINCIPAL -------------------------

def main():
    resultados= []
    config_path = os.environ.get('CONFIG_JSON_PATH', 'configs/linkedin.json')
    genai =  carrega_chave() # Use variável de ambiente para segurança

    try:
        # Carregar informações das configurações
        with open(config_path, "r", encoding="utf-8") as file:
            linkedin_config = json.load(file)
    except Exception as e:
        logger.error(f"Erro ao carregar configuração: {e}", file=sys.stderr)
        sys.exit(1)

    diretorio_cv=linkedin_config["input_file_cv"]
    output_dir = linkedin_config['output_dir']

    # Lê as vagas do JSON no n8n como a entrada é via stdin a linha de baixo não é necessaria 
    try:
        if stdin_has_data():
            entrada_json = sys.stdin.read().strip()
            logger.info("Lido JSON do stdin", file=sys.stderr)
        elif len(sys.argv) > 1:
            with open(sys.argv[1]) as f:
                entrada_json = f.read()
            logger.info(f"Lido JSON do arquivo {sys.argv[1]}", file=sys.stderr)
        else:
            logger.info("Nenhuma entrada fornecida! Use argumento de arquivo ou STDIN.", file=sys.stderr)
            sys.exit(1)
        vagas = json.loads(entrada_json)
    except Exception as e:
        logger.info(f"Erro ao ler entrada: {e}")
        sys.exit(1)
        

    for vaga in vagas:

        resultado=processar_vaga(genai, vaga, output_dir,diretorio_cv) 
        resultados.append(resultado)   

    print(json.dumps(resultados, ensure_ascii=False, indent=2))

    # O print dos logs para stderr foi REMOVIDO. O módulo logging já está fazendo isso em tempo real.
    logger.info("Processo finalizado.")




if __name__ == "__main__":
    main()